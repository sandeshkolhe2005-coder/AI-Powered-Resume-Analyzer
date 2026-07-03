import os
import pypdf
import docx
from typing import Optional

class ResumeParserService:
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        text = ""
        try:
            with open(file_path, "rb") as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise RuntimeError(f"Failed to read PDF file: {str(e)}")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]
            # Include text inside tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs.append(cell.text)
            return "\n".join(paragraphs)
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX file: {str(e)}")

    @classmethod
    def parse(cls, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            parsed_text = cls.extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            parsed_text = cls.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file extension. Only PDF and DOCX files are allowed.")
            
        # Clean extra whitespace
        cleaned_text = "\n".join([line.strip() for line in parsed_text.splitlines() if line.strip()])
        return cleaned_text
